# 负责把教师上传的各种格式文件PDF/Word/TXT/Markdown转换成纯文本，再切成小段chunk供 AI 使用
# python的PDF读取库
import PyPDF2
# python的Word读取库
from docx import Document as DocxDocument
# io.BytesIO = 字节数组转换为文件对象，用于读取二进制数据
import io 

def parse_pdf(file_content: bytes) -> str:
    # PDF2需要一个文件对象，所以用 io.BytesIO(file_content) 转换为文件对象，伪文件对象，保存在内存中，程序关闭消失
    # PyPDF2.PdfReader = PDF读取器，用于读取PDF文件
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        # page.extract_text() = 提取当前页的文本内容
        text += page.extract_text()
    return text

def parse_docx(file_content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_content))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
        # 读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text += " | ".join(row_text) + "\n"
    return text

def parse_txt_safe(file_content: bytes) -> str:
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文件编码")

def parse_markdown(file_content: bytes) -> str:
    return file_content.decode("utf-8")

def get_text_from_file(file_content: bytes, filename: str) -> str:
    ext = filename.split(".")[-1].lower() # 分解出文件后缀，转换成小写，实现兼容
    if ext == "pdf":
        return parse_pdf(file_content)
    elif ext == "docx":
        return parse_docx(file_content)
    elif ext == "txt":
        return parse_txt(file_content)
    elif ext == "md":
        return parse_markdown(file_content)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - chunk_overlap # 重叠的作用：确保跨边界的关键信息不会被切断，提高 RAG 检索的准确性
    return chunks
